from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Inches(0.5)

def read_file(filepath, lines=None):
    try:
        with open(filepath, 'r') as f:
            content = f.read().splitlines()
            if lines:
                return '\n'.join(content[:lines])
            return '\n'.join(content)
    except:
        return ""

doc = Document()

# Title
title = doc.add_heading('MLOps Assignment 2: Comprehensive Technical Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('This document provides an exhaustive, step-by-step technical breakdown of the end-to-end MLOps pipeline implemented for the binary image classification assignment.')
doc.add_page_break()

# Executive Overview & Findings
doc.add_heading('Project Overview', level=1)
doc.add_paragraph('This project implements an end-to-end MLOps pipeline for a binary image classification task (Cats vs Dogs) designed for a pet adoption platform. It covers data acquisition, Exploratory Data Analysis (EDA), model training with MLflow tracking, containerization, and CI/CD into a Kubernetes cluster.')

doc.add_heading('EDA Findings', level=1)
doc.add_paragraph('The dataset is perfectly balanced (~50% Cats, ~50% Dogs). The images have massive variance in lighting, background, and object scale. Preprocessing required standardized reshaping (224x224), normalization, and data augmentation to prevent overfitting to background noise.')

doc.add_heading('Model Comparison', level=1)
doc.add_paragraph('1. Baseline Custom CNN: Extremely fast inference (~5ms), tiny model size (<5MB), but lower accuracy (~70%).\n2. ResNet18 (Transfer Learning) (Selected): Vastly superior accuracy (~95%+), slightly slower inference but perfectly acceptable for REST API deployment.')

doc.add_heading('Setup Instructions & Repository', level=1)
doc.add_paragraph('Repository Link: [ INSERT YOUR REPOSITORY LINK HERE ]\n\nSetup requires a Python environment. Due to macOS vs Linux dependency constraints, two requirements files were utilized: requirements-mac.txt (for local Jupyter training) and requirements.txt (for GitHub CI/CD and Docker builds).')

doc.add_heading('Architecture & Screenshots', level=1)
doc.add_paragraph('[ INSERT ARCHITECTURE DIAGRAM SCREENSHOT HERE ]')
doc.add_paragraph('[ INSERT MLFLOW EXPERIMENT & ARTIFACT SCREENSHOTS HERE ]')
doc.add_paragraph('[ INSERT GITHUB ACTIONS CI/CD SCREENSHOTS HERE ]')
doc.add_paragraph('[ INSERT STREAMLIT UI & FASTAPI DEPLOYMENT SCREENSHOTS HERE ]')
doc.add_page_break()

# M1
doc.add_heading('Module 1: Containerization and API Enhancement', level=1)
doc.add_paragraph('The foundational step involved exposing the PyTorch ResNet-18 model via a REST API and containerizing it.')
doc.add_heading('FastAPI Backend', level=2)
doc.add_paragraph('FastAPI was used to serve the model. The /predict endpoint processes incoming multipart image uploads, passes them through torchvision transforms, and returns the inference result.')
add_code_block(doc, read_file('app/main.py', 25))

doc.add_heading('Dependency Management', level=2)
doc.add_paragraph('Strict versioning was enforced in requirements.txt to prevent conflicts between torchvision, streamlit, and Pillow.')
add_code_block(doc, read_file('requirements.txt'))

doc.add_heading('Dockerization', level=2)
doc.add_paragraph('A multi-stage optimized Dockerfile was created using python:3.9-slim to minimize the image footprint and maximize layer caching.')
add_code_block(doc, read_file('Dockerfile'))
doc.add_paragraph('[ INSERT FASTAPI SWAGGER DOCS SCREENSHOT HERE ]')
doc.add_page_break()

# M2
doc.add_heading('Module 2: Continuous Integration (CI)', level=1)
doc.add_paragraph('To ensure code quality, a fully automated CI pipeline was configured using GitHub Actions.')
doc.add_heading('Automated Testing & Build', level=2)
doc.add_paragraph('The CI job runs on ubuntu-latest. It sets up Python, installs dependencies, and runs Pytest. Upon success, it builds the Docker image and tags it dynamically using the Git SHA.')
doc.add_heading('DockerHub Push', level=2)
doc.add_paragraph('The pipeline securely authenticates with DockerHub using GitHub Secrets and pushes the image.')
add_code_block(doc, read_file('.github/workflows/ci-cd.yml', 40))
doc.add_paragraph('[ INSERT GITHUB ACTIONS CI PIPELINE SCREENSHOT HERE ]')
doc.add_page_break()

# M3 & M4
doc.add_heading('Module 3 & 4: Continuous Deployment (CD) and Testing', level=1)
doc.add_paragraph('The deployment strategy utilized Minikube for a local Kubernetes environment.')
doc.add_heading('Kubernetes Manifests', level=2)
doc.add_paragraph('Declarative YAML manifests were authored for the backend and frontend. The Streamlit UI routes traffic to the backend via Cluster DNS (http://mlops-inference-service:8000).')
add_code_block(doc, read_file('k8s/deployment.yaml', 20))

doc.add_heading('Self-Hosted Runner & Local Deploy Script', level=2)
doc.add_paragraph('To deploy directly into the local macOS Minikube cluster, a Self-Hosted GitHub Runner was configured. The deployment script dynamically rewrites the YAML files with the latest Git SHA tag.')
add_code_block(doc, read_file('scripts/deploy_local.sh', 30))

doc.add_heading('Automated Smoke Testing', level=2)
doc.add_paragraph('To bypass macOS Minikube networking limitations, a background port-forward tunnel is established before running the smoke tests to ensure the live API is responsive.')
add_code_block(doc, read_file('scripts/smoke_test.py', 25))
doc.add_paragraph('[ INSERT TERMINAL SCREENSHOT SHOWING PODS RUNNING HERE ]')
doc.add_page_break()

# M5
doc.add_heading('Module 5: Monitoring, Logs, and Observability', level=1)
doc.add_paragraph('Real-time monitoring and post-deployment performance tracking were integrated into the pipeline.')
doc.add_heading('Prometheus Metrics', level=2)
doc.add_paragraph('The prometheus-fastapi-instrumentator library was used to expose a /metrics endpoint, capturing CPU time, Memory, Request Counts, and Latency.')

doc.add_heading('Streamlit Monitoring Dashboard', level=2)
doc.add_paragraph('A custom dashboard was engineered into Streamlit. It parses the Prometheus scientific notation output to dynamically visualize API Health.')
add_code_block(doc, read_file('app/ui.py', 65))

doc.add_heading('Live Performance Benchmark (Data Drift Simulation)', level=2)
doc.add_paragraph('A Performance Tracking tool pulls real, pre-labeled images (bundled in app/assets), sends them to the live API, and calculates the overall live Model Accuracy to detect potential data drift.')
doc.add_paragraph('[ INSERT STREAMLIT MONITORING DASHBOARD SCREENSHOT HERE ]')

doc.save('docs/Assignment_Report.docx')
print("Successfully generated ultra-detailed docs/Assignment_Report.docx")
